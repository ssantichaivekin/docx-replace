from docx import Document
from docx.shared import Pt
import json
import re

## replace text in runs :

def applyOptions(run, font_options = []) :
    '''
    Aply options to a run. Current options:
    You can look at option definitions below.
    '''
    # define options:
    def decreaseSizeOption(run, max_len, size_to_decrease) :
        # If the replaced text is long, decrease the font
        old_size = run.font.size.pt
        if len(run.text) > max_len :
            new_size = old_size - Pt(size_to_decrease)
            run.font.size = new_size
    
    # apply options:
    for option in font_options :
        if option['name'] == 'decrease_size' :
            decreaseSizeOption(run, *option[args])
            

def replaceRun(run, old_text, new_text, font_options = []) :
    '''
    Replace occurances of textA with textB in a run in place,
    Font options can be specified.
    '''
    # special option :
    old_run_text = run.text
    new_run_text = old_run_text.replace(old_text, new_text)
    if old_run_text != new_run_text :
        run.text = new_run_text
        applyOptions(run, font_options)

def replaceRunWithRegex(run, pattern, repl, font_options = []) :
    old_run_text = run.text
    new_run_text = re.sub(pattern, repl, old_run_text)
    if old_run_text != new_run_text :
        run.text = new_run_text
        applyOptions(run, font_options)

def enumerateRun(document) :
    for para in document.paragraphs :
        for run in para.runs :
            yield run

    # replace in tables
    for table in document.tables :
        for row in table.rows :
            for cell in row.cells :
                for para in cell.paragraphs :
                    for run in para.runs :
                        yield run

def replace(document, old_text, new_text, font_options = []) :
    '''
    Replace occurances of textA with textB in docx document in place,
    preserving original style of textA. Font options can be specified
    to change the style of the newly replaced textB.
    '''
    for run in enumerateRun(document) :
        replaceRun(run, old_text, new_text, font_options)

def replaceWithRegex(document, pattern, repl, font_options = []) :
    for run in enumerateRun(document) :
        replaceRunWithRegex(run, pattern, repl, font_options)

def multiReplace(document, replacement_list, global_options) :
    '''
    Replace using multiple old text and new text from the list.
    You can also specify more global options using the global_options dict.
    Possible options: use_braces = ['a', 'b'], remove_empty_braces = True
    '''
    # parse global options :
    use_braces = 'use_braces' in global_options and len(global_options['use_braces']) == 2
    if use_braces :
        open_brace = global_options['use_braces'][0]
        close_brace = global_options['use_braces'][1]
    remove_empty_braces = 'remove_empty_braces' in global_options and global_options['remove_empty_braces']
    remove_unreplaced_braces = 'remove_unreplaced_braces' in global_options and global_options['remove_unreplaced_braces']

    # do the work :
    for replacement in replacement_list:
        old_text = replacement['old_text']
        new_text = replacement['new_text']
        if 'options' in replacement :
            options = replacement['options']
        else :
            options = []
        if remove_unreplaced_braces :
            options += []
        if use_braces :
            old_text = open_brace + old_text + close_brace
        replace(document, old_text, new_text, options)

    if remove_empty_braces :
        old_text = open_brace + close_brace
        replace(document, old_text, '')
    
    if remove_unreplaced_braces :
        re_open_brace = re.escape(open_brace)
        re_close_brace = re.escape(close_brace)
        # maps anything that is open_brace...close_brace (non greedy)
        re_pattern = '%s.*?%s' % (re_open_brace, re_close_brace)
        replaceWithRegex(document, re_pattern, '')

## read / write

def readReplaceListFromJson(jsonfilepath) :
    with open(jsonfilepath, 'r') as jsonfile :
        replacement_list = json.load(jsonfile)
        return replacement_list

## demo

def demo() :
    print('Replace demo -- replace () to [] in demo.docx', end=' ')
    print('save the new document as demo2.docx')
    document = Document('demo.docx')
    replace(document, '(', '[')
    replace(document, ')', ']')
    document.save('demo2.docx')

if __name__ == '__main__' :
    import argparse
    parser = argparse.ArgumentParser(description='Replace text in docx')
    parser.add_argument('original_docx_path')
    parser.add_argument('old_text', nargs='?', default=None)
    parser.add_argument('new_text', nargs='?', default=None)
    parser.add_argument('--replace-list-file')
    parser.add_argument('--dest')
    parser.add_argument('--use-braces', nargs=2, default=None)
    parser.add_argument('--remove-empty-braces', action='store_true')
    parser.add_argument('--remove-unreplaced-braces', action='store_true')
    parser.add_argument('--remove-empty-row', type=int)

    args = parser.parse_args()
    
    # input validity checks
    if not args.old_text and not args.replace_list_file :
        raise Exception('You must specify either old_text and new_text argument or use a json spec file')
    
    if args.old_text and not args.new_text :
        raise Exception('You must specify new_text if you specified old_text', args.old_text, args.new_text)

    # process
    document = Document(args.original_docx_path)
    if args.replace_list_file :
        replace_list = readReplaceListFromJson(args.replace_list_file)
    else :
        replace_list = [{ "old_text": args.old_text, "new_text": args.new_text }]
    
    multiReplace(document, replace_list, vars(args))
    
    # save
    dest = args.dest or args.original_docx_path
    document.save(dest)
    



    


